import re
from io import BytesIO
from flask import Blueprint, render_template, request, send_file
from flask_login import login_required
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from docx import Document as DocxDocument
from app import db, Production, CastEntry, Artist

reports_bp = Blueprint('reports', __name__, url_prefix='/reports')

XLSX_MIME = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
DOCX_MIME = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

def safe_filename(name):
    return re.sub(r'[\\/:*?"<>|]', '_', name).strip() or 'report'

def xlsx_response(headers, rows, filename):
    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for row in rows:
        ws.append(row)
    for i, h in enumerate(headers, 1):
        ws.column_dimensions[get_column_letter(i)].width = max(14, len(str(h)) + 4)
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name=f'{safe_filename(filename)}.xlsx', mimetype=XLSX_MIME)

def docx_table_response(title, headers, rows, filename):
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val) if val not in (None, '') else '—'
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name=f'{safe_filename(filename)}.docx', mimetype=DOCX_MIME)

@reports_bp.route('/')
@login_required
def index():
    genres = sorted({p.genre for p in Production.query.all() if p.genre})
    return render_template('reports/index.html', genres=genres)

# ---------- Артисты: участие в постановках за период ----------
@reports_bp.route('/artists/export')
@login_required
def export_artists():
    fmt = request.args.get('format', 'xlsx')
    q = request.args.get('q', '').strip()
    year_from = request.args.get('year_from', '').strip()
    year_to = request.args.get('year_to', '').strip()

    query = CastEntry.query.join(Artist).join(Production)
    if q:
        query = query.filter(Artist.full_name.ilike(f'%{q}%'))
    entries = query.order_by(Artist.full_name, Production.premiere_year).all()

    yf = int(year_from) if year_from else None
    yt = int(year_to) if year_to else None

    def overlaps(ce):
        if not yf and not yt:
            return True
        ce_from = ce.year_from or ce.production.premiere_year
        ce_to = ce.year_to or ce_from
        if yf and ce_to < yf:
            return False
        if yt and ce_from > yt:
            return False
        return True

    headers = ['ФИО артиста', 'Звание', 'Постановка', 'Роль', 'Годы участия']
    rows = [
        [ce.artist.full_name, ce.artist.title or '', ce.production.name, ce.role_name or '', ce.years_display or '']
        for ce in entries if overlaps(ce)
    ]

    title = 'Отчёт по артистам'
    if yf or yt:
        title += f' за {yf or "..."}–{yt or "..."}'
    if fmt == 'docx':
        return docx_table_response(title, headers, rows, title)
    return xlsx_response(headers, rows, title)

# ---------- Постановки: фильтруемый список ----------
@reports_bp.route('/productions/export')
@login_required
def export_productions():
    fmt = request.args.get('format', 'xlsx')
    status = request.args.get('status', '')
    genre = request.args.get('genre', '').strip()
    year_from = request.args.get('year_from', '').strip()
    year_to = request.args.get('year_to', '').strip()
    basis = request.args.get('basis', '').strip()

    query = Production.query
    if status:
        query = query.filter(Production.status == status)
    if genre:
        query = query.filter(Production.genre.ilike(f'%{genre}%'))
    if year_from:
        query = query.filter(Production.premiere_year >= int(year_from))
    if year_to:
        query = query.filter(Production.premiere_year <= int(year_to))
    if basis:
        query = query.filter(db.or_(
            Production.literary_basis.ilike(f'%{basis}%'),
            Production.literary_basis_author.ilike(f'%{basis}%'),
        ))
    productions = query.order_by(Production.premiere_year.desc()).all()

    headers = ['Название', 'Статус', 'Жанр', 'Премьера', 'Актов',
               'Авторы музыки', 'Авторы либретто', 'Литературная основа', 'Автор основы']
    rows = [[
        p.name, p.status_display, p.genre, p.premiere_display, p.acts_count or '',
        p.music_authors_display, p.libretto_authors_display,
        p.literary_basis or '', p.literary_basis_author or '',
    ] for p in productions]

    title = 'Отчёт по постановкам'
    if fmt == 'docx':
        return docx_table_response(title, headers, rows, title)
    return xlsx_response(headers, rows, title)

# ---------- Полная информация по одной постановке ----------
def _production_docx(p):
    doc = DocxDocument()
    doc.add_heading(p.name, level=1)
    doc.add_paragraph(f'Статус: {p.status_display}')
    doc.add_paragraph(f'Жанр: {p.genre}')
    doc.add_paragraph(f'Премьера: {p.premiere_display}')
    if p.acts_count:
        doc.add_paragraph(f'Количество актов: {p.acts_count}')
    if p.music_authors:
        doc.add_paragraph(f'Авторы музыки: {p.music_authors_display}')
    if p.libretto_authors:
        doc.add_paragraph(f'Авторы либретто: {p.libretto_authors_display}')
    if p.literary_basis:
        line = f'Литературная основа: {p.literary_basis}'
        if p.literary_basis_author:
            line += f' (автор: {p.literary_basis_author})'
        doc.add_paragraph(line)

    doc.add_heading('Состав исполнителей', level=2)
    if p.cast_entries:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = 'Артист', 'Роль', 'Годы участия'
        for ce in p.cast_entries:
            cells = table.add_row().cells
            cells[0].text = ce.artist.full_name
            cells[1].text = ce.role_name or '—'
            cells[2].text = ce.years_display or '—'
    else:
        doc.add_paragraph('Нет данных.')

    doc.add_heading('Постановочная группа', level=2)
    if p.staging_group:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = 'ФИО', 'Должность', 'Годы жизни'
        for pd in p.staging_group:
            cells = table.add_row().cells
            cells[0].text = pd.director.full_name
            cells[1].text = pd.director.position_display
            cells[2].text = pd.director.life_years or '—'
    else:
        doc.add_paragraph('Нет данных.')

    doc.add_heading('Документы', level=2)
    if p.documents:
        for d in p.documents:
            doc.add_paragraph(f'{d.doc_type_display}: {d.title or d.file_name}', style='List Bullet')
    else:
        doc.add_paragraph('Нет данных.')

    doc.add_heading('Материалы', level=2)
    if p.materials:
        for m in p.materials:
            doc.add_paragraph(f'{m.type_display}: {m.title or m.file_name or m.url}', style='List Bullet')
    else:
        doc.add_paragraph('Нет данных.')

    return doc

def _production_xlsx(p):
    wb = Workbook()
    ws = wb.active
    ws.title = 'Информация'
    ws.append(['Поле', 'Значение'])
    ws.append(['Название', p.name])
    ws.append(['Статус', p.status_display])
    ws.append(['Жанр', p.genre])
    ws.append(['Премьера', p.premiere_display])
    ws.append(['Актов', p.acts_count or ''])
    ws.append(['Авторы музыки', p.music_authors_display])
    ws.append(['Авторы либретто', p.libretto_authors_display])
    ws.append(['Литературная основа', p.literary_basis or ''])
    ws.append(['Автор основы', p.literary_basis_author or ''])
    for col, width in (('A', 22), ('B', 50)):
        ws.column_dimensions[col].width = width

    ws2 = wb.create_sheet('Состав')
    ws2.append(['Артист', 'Роль', 'Годы участия'])
    for ce in p.cast_entries:
        ws2.append([ce.artist.full_name, ce.role_name or '', ce.years_display or ''])

    ws3 = wb.create_sheet('Постановочная группа')
    ws3.append(['ФИО', 'Должность', 'Годы жизни'])
    for pd in p.staging_group:
        ws3.append([pd.director.full_name, pd.director.position_display, pd.director.life_years or ''])

    ws4 = wb.create_sheet('Документы')
    ws4.append(['Тип', 'Название', 'Файл'])
    for d in p.documents:
        ws4.append([d.doc_type_display, d.title or '', d.file_name])

    ws5 = wb.create_sheet('Материалы')
    ws5.append(['Тип', 'Название', 'Файл/ссылка'])
    for m in p.materials:
        ws5.append([m.type_display, m.title or '', m.file_name or m.url or ''])

    for sheet in (ws2, ws3, ws4, ws5):
        for i, _ in enumerate(sheet[1], 1):
            sheet.column_dimensions[get_column_letter(i)].width = 28

    return wb

@reports_bp.route('/production/<int:pid>/export')
@login_required
def export_production(pid):
    p = db.get_or_404(Production, pid)
    fmt = request.args.get('format', 'docx')
    filename = safe_filename(p.name)
    if fmt == 'xlsx':
        wb = _production_xlsx(p)
        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True, download_name=f'{filename}.xlsx', mimetype=XLSX_MIME)
    doc = _production_docx(p)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name=f'{filename}.docx', mimetype=DOCX_MIME)
